"""
v18.0 — Document & Content Management
Rich text editor, templates, version history, collaborative editing, export
"""
import json
import logging
import hashlib
from fastapi import APIRouter, Depends, HTTPException, Header
from pydantic import BaseModel
from typing import Optional
from config import settings as app_settings
import db.client as db

logger = logging.getLogger(__name__)
router = APIRouter()

def verify_key(x_api_key: Optional[str] = Header(None)):
    if x_api_key != app_settings.dawn_api_key:
        raise HTTPException(status_code=401, detail="Invalid API key")

# ─── Schemas ──────────────────────────────────────────────────────────────

class DocumentCreate(BaseModel):
    title: str
    content: str = ""
    content_type: str = "markdown"  # 'markdown', 'html', 'richtext'
    folder_id: Optional[str] = None
    tags: list[str] = []
    template_id: Optional[str] = None

class DocumentUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    content_type: Optional[str] = None
    folder_id: Optional[str] = None
    tags: Optional[list[str]] = None

class FolderCreate(BaseModel):
    name: str
    parent_id: Optional[str] = None
    description: str = ""

class TemplateCreate(BaseModel):
    name: str
    description: str = ""
    content: str = ""
    content_type: str = "markdown"
    variables: list[dict] = []  # [{"name": "title", "type": "string", "default": ""}]

# ─── Documents CRUD ───────────────────────────────────────────────────────

@router.get("/documents", tags=["documents"])
async def list_documents(
    folder_id: Optional[str] = None,
    tag: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
    _: None = Depends(verify_key),
):
    """List documents with optional filtering."""
    try:
        supabase = db.get_db()
        q = supabase.table("documents").select("*").order("updated_at", desc=True)
        
        if folder_id:
            q = q.eq("folder_id", folder_id)
        if tag:
            q = q.contains("tags", [tag])
        
        res = q.limit(limit).offset(offset).execute()
        return res.data or []
    except Exception as e:
        logger.error(f"[documents] list failed: {e}")
        return []


@router.get("/documents/{doc_id}", tags=["documents"])
async def get_document(doc_id: str, _: None = Depends(verify_key)):
    """Get a document by ID."""
    try:
        supabase = db.get_db()
        res = supabase.table("documents").select("*").eq("id", doc_id).execute()
        if not res.data:
            raise HTTPException(status_code=404, detail="Document not found")
        return res.data[0]
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[documents] get failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get document: {str(e)}")


@router.post("/documents", tags=["documents"])
async def create_document(req: DocumentCreate, _: None = Depends(verify_key)):
    """Create a new document."""
    try:
        supabase = db.get_db()
        
        data = req.model_dump()
        data["content_hash"] = hashlib.sha256(req.content.encode()).hexdigest()[:16]
        
        res = supabase.table("documents").insert(data).execute()
        if not res.data:
            raise HTTPException(status_code=500, detail="Failed to create document")
        
        # Create initial version
        doc = res.data[0]
        supabase.table("document_versions").insert({
            "document_id": doc["id"],
            "content": req.content,
            "version_number": 1,
            "change_note": "Initial version",
            "content_hash": data["content_hash"],
        }).execute()
        
        return doc
    except Exception as e:
        logger.error(f"[documents] create failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create document: {str(e)}")


@router.put("/documents/{doc_id}", tags=["documents"])
async def update_document(
    doc_id: str,
    req: DocumentUpdate,
    change_note: str = "",
    _: None = Depends(verify_key),
):
    """Update a document and create a new version if content changed."""
    try:
        supabase = db.get_db()
        
        # Get current document
        current = supabase.table("documents").select("*").eq("id", doc_id).execute()
        if not current.data:
            raise HTTPException(status_code=404, detail="Document not found")
        
        current_doc = current.data[0]
        update_data = {k: v for k, v in req.model_dump().items() if v is not None}
        
        # If content changed, create new version
        if req.content is not None and req.content != current_doc.get("content", ""):
            new_hash = hashlib.sha256(req.content.encode()).hexdigest()[:16]
            update_data["content_hash"] = new_hash
            update_data["version"] = current_doc.get("version", 1) + 1
            
            # Get next version number
            versions = supabase.table("document_versions").select("version_number").eq(
                "document_id", doc_id
            ).order("version_number", desc=True).limit(1).execute()
            
            next_version = (versions.data[0]["version_number"] + 1) if versions.data else 1
            
            supabase.table("document_versions").insert({
                "document_id": doc_id,
                "content": req.content,
                "version_number": next_version,
                "change_note": change_note or f"Version {next_version}",
                "content_hash": new_hash,
            }).execute()
        
        update_data["updated_at"] = "now()"
        res = supabase.table("documents").update(update_data).eq("id", doc_id).execute()
        
        return res.data[0] if res.data else current_doc
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[documents] update failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to update document: {str(e)}")


@router.delete("/documents/{doc_id}", tags=["documents"])
async def delete_document(doc_id: str, _: None = Depends(verify_key)):
    """Delete a document."""
    try:
        supabase = db.get_db()
        supabase.table("documents").delete().eq("id", doc_id).execute()
        return {"status": "deleted"}
    except Exception as e:
        logger.error(f"[documents] delete failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")


# ─── Document Versions ────────────────────────────────────────────────────

@router.get("/documents/{doc_id}/versions", tags=["documents"])
async def list_document_versions(doc_id: str, _: None = Depends(verify_key)):
    """List all versions of a document."""
    try:
        supabase = db.get_db()
        res = supabase.table("document_versions").select(
            "id, version_number, change_note, content_hash, created_at"
        ).eq("document_id", doc_id).order("version_number", desc=True).execute()
        return res.data or []
    except Exception as e:
        logger.error(f"[documents] list versions failed: {e}")
        return []


@router.get("/documents/{doc_id}/versions/{version_id}", tags=["documents"])
async def get_document_version(
    doc_id: str,
    version_id: str,
    _: None = Depends(verify_key),
):
    """Get a specific version of a document."""
    try:
        supabase = db.get_db()
        res = supabase.table("document_versions").select("*").eq(
            "id", version_id
        ).eq("document_id", doc_id).execute()
        if not res.data:
            raise HTTPException(status_code=404, detail="Version not found")
        return res.data[0]
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[documents] get version failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get version: {str(e)}")


@router.post("/documents/{doc_id}/restore/{version_id}", tags=["documents"])
async def restore_document_version(
    doc_id: str,
    version_id: str,
    _: None = Depends(verify_key),
):
    """Restore a document to a previous version."""
    try:
        supabase = db.get_db()
        
        # Get the version
        version = supabase.table("document_versions").select("*").eq(
            "id", version_id
        ).eq("document_id", doc_id).execute()
        
        if not version.data:
            raise HTTPException(status_code=404, detail="Version not found")
        
        # Update document with version content
        version_data = version.data[0]
        new_hash = hashlib.sha256(version_data["content"].encode()).hexdigest()[:16]
        
        supabase.table("documents").update({
            "content": version_data["content"],
            "content_hash": new_hash,
            "version": supabase.table("documents").select("version").eq("id", doc_id).execute().data[0].get("version", 1) + 1,
        }).eq("id", doc_id).execute()
        
        return {"status": "restored", "version": version_data["version_number"]}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[documents] restore failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to restore version: {str(e)}")


# ─── Folders ──────────────────────────────────────────────────────────────

@router.get("/documents/folders", tags=["documents"])
async def list_folders(parent_id: Optional[str] = None, _: None = Depends(verify_key)):
    """List document folders."""
    try:
        supabase = db.get_db()
        q = supabase.table("document_folders").select("*").order("name")
        if parent_id:
            q = q.eq("parent_id", parent_id)
        else:
            q = q.is_("parent_id", "null")
        res = q.execute()
        return res.data or []
    except Exception as e:
        logger.error(f"[documents] list folders failed: {e}")
        return []


@router.post("/documents/folders", tags=["documents"])
async def create_folder(req: FolderCreate, _: None = Depends(verify_key)):
    """Create a document folder."""
    try:
        supabase = db.get_db()
        res = supabase.table("document_folders").insert(req.model_dump()).execute()
        return res.data[0] if res.data else {}
    except Exception as e:
        logger.error(f"[documents] create folder failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create folder: {str(e)}")


@router.delete("/documents/folders/{folder_id}", tags=["documents"])
async def delete_folder(folder_id: str, _: None = Depends(verify_key)):
    """Delete a folder."""
    try:
        supabase = db.get_db()
        supabase.table("document_folders").delete().eq("id", folder_id).execute()
        return {"status": "deleted"}
    except Exception as e:
        logger.error(f"[documents] delete folder failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete folder: {str(e)}")


# ─── Templates ────────────────────────────────────────────────────────────

@router.get("/documents/templates", tags=["documents"])
async def list_templates(_: None = Depends(verify_key)):
    """List document templates."""
    try:
        supabase = db.get_db()
        res = supabase.table("document_templates").select("*").order("name").execute()
        return res.data or []
    except Exception as e:
        logger.error(f"[documents] list templates failed: {e}")
        return []


@router.post("/documents/templates", tags=["documents"])
async def create_template(req: TemplateCreate, _: None = Depends(verify_key)):
    """Create a document template."""
    try:
        supabase = db.get_db()
        res = supabase.table("document_templates").insert(req.model_dump()).execute()
        return res.data[0] if res.data else {}
    except Exception as e:
        logger.error(f"[documents] create template failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create template: {str(e)}")


@router.post("/documents/templates/{template_id}/apply", tags=["documents"])
async def apply_template(
    template_id: str,
    variables: dict = {},
    _: None = Depends(verify_key),
):
    """Apply a template with variable substitution and create a new document."""
    try:
        supabase = db.get_db()
        
        # Get template
        template = supabase.table("document_templates").select("*").eq("id", template_id).execute()
        if not template.data:
            raise HTTPException(status_code=404, detail="Template not found")
        
        tmpl = template.data[0]
        content = tmpl["content"]
        
        # Substitute variables
        for key, value in variables.items():
            content = content.replace(f"{{{{ {key} }}}}", str(value))
            content = content.replace(f"{{{{{key}}}}}", str(value))
        
        # Create document from template
        doc = supabase.table("documents").insert({
            "title": variables.get("title", tmpl["name"]),
            "content": content,
            "content_type": tmpl["content_type"],
            "tags": [],
        }).execute()
        
        return doc.data[0] if doc.data else {}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[documents] apply template failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to apply template: {str(e)}")


# ─── Export ───────────────────────────────────────────────────────────────

@router.get("/documents/{doc_id}/export", tags=["documents"])
async def export_document(
    doc_id: str,
    format: str = "markdown",  # 'markdown', 'html', 'pdf', 'docx'
    _: None = Depends(verify_key),
):
    """Export a document in various formats."""
    try:
        supabase = db.get_db()
        res = supabase.table("documents").select("*").eq("id", doc_id).execute()
        if not res.data:
            raise HTTPException(status_code=404, detail="Document not found")
        
        doc = res.data[0]
        content = doc.get("content", "")
        title = doc.get("title", "Untitled")
        
        if format == "markdown":
            return {"content": content, "format": "markdown", "filename": f"{title}.md"}
        
        elif format == "html":
            import markdown
            html = markdown.markdown(content, extensions=["fenced_code", "tables", "codehilite"])
            full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><title>{title}</title>
<style>body{{max-width:800px;margin:auto;padding:2em;font-family:system-ui,sans-serif;line-height:1.6}}
code{{background:#f4f4f4;padding:2px 6px;border-radius:3px}}
pre code{{display:block;padding:1em;overflow-x:auto}}
table{{border-collapse:collapse;width:100%}}
th,td{{border:1px solid #ddd;padding:8px;text-align:left}}
</style></head><body>{html}</body></html>"""
            return {"content": full_html, "format": "html", "filename": f"{title}.html"}
        
        elif format == "pdf":
            try:
                from weasyprint import HTML
                import markdown
                import tempfile
                
                html = markdown.markdown(content)
                full_html = f"<h1>{title}</h1>{html}"
                
                with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                    HTML(string=full_html).write_pdf(tmp.name)
                    with open(tmp.name, "rb") as f:
                        import base64
                        pdf_base64 = base64.b64encode(f.read()).decode()
                    import os
                    os.unlink(tmp.name)
                
                return {"content": pdf_base64, "format": "pdf", "filename": f"{title}.pdf", "encoded": True}
            except ImportError:
                raise HTTPException(status_code=501, detail="PDF export requires weasyprint: pip install weasyprint")
        
        elif format == "docx":
            try:
                from docx import Document
                import markdown
                from markdown_to_docx import convert
                import tempfile
                import base64
                
                docx = Document()
                docx.add_heading(title, 0)
                docx.add_paragraph(content)
                
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    docx.save(tmp.name)
                    with open(tmp.name, "rb") as f:
                        docx_base64 = base64.b64encode(f.read()).decode()
                    import os
                    os.unlink(tmp.name)
                
                return {"content": docx_base64, "format": "docx", "filename": f"{title}.docx", "encoded": True}
            except ImportError:
                raise HTTPException(status_code=501, detail="DOCX export requires python-docx: pip install python-docx")
        
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"[documents] export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


# ─── Search ───────────────────────────────────────────────────────────────

@router.get("/documents/search", tags=["documents"])
async def search_documents(
    q: str,
    limit: int = 20,
    _: None = Depends(verify_key),
):
    """Search documents by title and content."""
    try:
        supabase = db.get_db()
        # Use PostgreSQL full-text search
        res = supabase.rpc("search_documents", {
            "search_query": q,
            "max_results": limit,
        }).execute()
        return res.data or []
    except Exception:
        # Fallback to simple ILIKE search
        try:
            supabase = db.get_db()
            res = supabase.table("documents").select("*").or_(
                f"title.ilike.%{q}%,content.ilike.%{q}%"
            ).limit(limit).execute()
            return res.data or []
        except Exception as e:
            logger.error(f"[documents] search failed: {e}")
            return []
